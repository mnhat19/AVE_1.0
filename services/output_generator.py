from __future__ import annotations

import json
from datetime import datetime, timezone
from pathlib import Path

import openpyxl
from openpyxl.styles import Alignment, Font, PatternFill
from docx import Document

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.pdfgen import canvas
except Exception:
    letter = None
    canvas = None

from config.llm import get_llm_client
from config.settings import settings

OUTPUT_DIR = Path(settings.output_dir)
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

SEVERITY_COLORS = {
    "CRITICAL": "FF0000",
    "HIGH": "FF6600",
    "MEDIUM": "FFCC00",
    "LOW": "99CC00",
}

MEMO_GENERATION_PROMPT = """You are a senior auditor writing an audit memo in Vietnamese.

Session: {session_id}
Stage: {stage}
Findings (use exact finding_id; do not invent new IDs): {findings_json}
Anomaly Summary: {anomaly_summary}

Structure:
1. Audit objective
2. Summary of procedures performed
3. Key findings (prioritized)
4. Root cause and expected impact
5. Recommended next steps

Write in formal Vietnamese audit language. Be specific and reference the finding IDs exactly as provided."""


def _build_consolidated_maps(state: dict) -> tuple[dict[int, dict], dict[int, dict]]:
    interim_map: dict[int, dict] = {}
    fieldwork_map: dict[int, dict] = {}
    for entry in state.get("consolidated_findings", []) or []:
        interim_index = entry.get("interim_finding_index")
        fieldwork_index = entry.get("fieldwork_finding_index")
        if interim_index:
            interim_map[int(interim_index)] = entry
        if fieldwork_index:
            fieldwork_map[int(fieldwork_index)] = entry
    return interim_map, fieldwork_map


def _materiality_from_severity(severity: str) -> str:
    if severity in {"CRITICAL", "HIGH"}:
        return "HIGHLY_MATERIAL"
    if severity == "MEDIUM":
        return "MATERIAL"
    return "IMMATERIAL"


def _annotate_findings(state: dict) -> list[dict]:
    findings = state.get("audit_findings", []) or []
    interim_idx = 0
    fieldwork_idx = 0
    for idx, finding in enumerate(findings, 1):
        stage = finding.get("stage") or state.get("stage") or "BOTH"
        if stage == "INTERIM":
            interim_idx += 1
            stage_index = interim_idx
        elif stage == "FIELDWORK":
            fieldwork_idx += 1
            stage_index = fieldwork_idx
        else:
            stage_index = idx

        finding.setdefault("stage", stage)
        finding["finding_id"] = f"FND-{state['session_id'][:4].upper()}-{idx:03d}"
        finding["_stage_index"] = stage_index
    return findings


def _file_name_map(state: dict) -> dict[str, str]:
    mapping: dict[str, str] = {}
    for record in state.get("file_records", []) or []:
        file_id = record.get("file_id") or record.get("id")
        filename = record.get("filename")
        if file_id and filename:
            mapping[file_id] = filename
    return mapping


def _render_evidence_links(evidence_links: list[dict], file_names: dict[str, str]) -> str:
    rendered: list[str] = []
    for link in evidence_links:
        reference = link.get("reference")
        if not reference:
            continue
        filename = file_names.get(link.get("source_file_id"), "")
        rendered.append(f"{filename}: {reference}" if filename else reference)
    return "; ".join(rendered)


def _find_finding_index(state: dict, description: str) -> int | None:
    for idx, finding in enumerate(state.get("audit_findings", []), 1):
        if finding.get("description") == description:
            return idx
    return None


async def generate_issue_log(state: dict) -> str:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Issue Log"

    headers = [
        "Finding ID",
        "Stage",
        "Description",
        "Root Cause",
        "Expected Impact",
        "Severity",
        "Materiality",
        "Review Flag",
        "Assignee",
        "Status",
        "Confidence Score",
        "Evidence Reference",
    ]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1F4E79")
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
    ws.row_dimensions[1].height = 30

    session_id = state["session_id"]
    findings = _annotate_findings(state)
    interim_map, fieldwork_map = _build_consolidated_maps(state)
    file_names = _file_name_map(state)
    for idx, finding in enumerate(findings, 1):
        row_num = idx + 1
        finding_id = finding.get("finding_id") or f"FND-{session_id[:4].upper()}-{idx:03d}"
        severity = finding.get("severity", "MEDIUM")
        stage = finding.get("stage") or state.get("stage") or "BOTH"
        stage_index = int(finding.get("_stage_index") or idx)
        consolidated = (
            interim_map.get(stage_index) if stage == "INTERIM" else fieldwork_map.get(stage_index)
        ) or {}
        materiality = consolidated.get("materiality") or _materiality_from_severity(severity)
        review_needed = consolidated.get("review_flag") or finding.get("manual_review_required")
        review_flag = "YES" if review_needed else "NO"
        evidence_links = finding.get("evidence_links") or []
        evidence_reference = _render_evidence_links(evidence_links, file_names)
        if not evidence_reference:
            evidence_reference = (
                finding.get("source_reference")
                or ", ".join(finding.get("related_anomaly_rules", []))
            )

        row_data = [
            finding_id,
            stage,
            finding.get("description", ""),
            finding.get("root_cause", ""),
            finding.get("expected_impact", ""),
            severity,
            materiality or "",
            review_flag,
            "Audit Team",
            "OPEN",
            f"{finding.get('confidence_score', 0.0):.0%}",
            evidence_reference,
        ]
        for col, val in enumerate(row_data, 1):
            ws.cell(row=row_num, column=col, value=val)

        sev_cell = ws.cell(row=row_num, column=6)
        sev_cell.fill = PatternFill("solid", fgColor=SEVERITY_COLORS.get(severity, "CCCCCC"))
        sev_cell.font = Font(bold=True)

    col_widths = [15, 10, 60, 50, 50, 12, 16, 12, 15, 12, 18, 40]
    for col, width in enumerate(col_widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(col)].width = width

    ws2 = wb.create_sheet("Risk Register")
    risk_headers = [
        "Risk ID",
        "Description",
        "Probability",
        "Impact",
        "Risk Score",
        "Owner",
        "Controls",
        "Control Test Result",
        "Consolidated Ref",
    ]
    for col, header in enumerate(risk_headers, 1):
        ws2.cell(row=1, column=col, value=header).font = Font(bold=True)

    interim_map, fieldwork_map = _build_consolidated_maps(state)
    risk_matrix = state.get("risk_control_matrix", []) or []
    for idx, risk in enumerate(state["risk_entries"], 2):
        finding_index = risk.get("finding_stage_index") or risk.get("finding_index")
        if not finding_index:
            finding_index = _find_finding_index(state, risk.get("description", ""))
        finding_stage = risk.get("finding_stage") or state.get("stage") or "BOTH"
        consolidated = (
            interim_map.get(finding_index) if finding_stage == "INTERIM" else fieldwork_map.get(finding_index)
        ) or {}
        consolidated_ref = f"CF-{finding_index:03d}" if consolidated and finding_index else None

        control_result = ""
        related_controls = risk.get("related_controls", []) or []
        for control in related_controls:
            match = next(
                (
                    entry
                    for entry in risk_matrix
                    if control.lower() in str(entry.get("control_description", "")).lower()
                ),
                None,
            )
            if match:
                control_result = match.get("test_result") or ""
                break
        ws2.cell(row=idx, column=1, value=risk.get("id", f"RISK-{idx}"))
        ws2.cell(row=idx, column=2, value=risk.get("description", ""))
        ws2.cell(row=idx, column=3, value=risk.get("probability", 0))
        ws2.cell(row=idx, column=4, value=risk.get("impact", 0))
        ws2.cell(row=idx, column=5, value=risk.get("risk_score", 0))
        ws2.cell(row=idx, column=6, value=risk.get("owner", ""))
        ws2.cell(row=idx, column=7, value=str(risk.get("related_controls", [])))
        ws2.cell(row=idx, column=8, value=control_result)
        ws2.cell(row=idx, column=9, value=consolidated_ref)

    path = OUTPUT_DIR / f"issue_log_{session_id}.xlsx"
    wb.save(path)
    return str(path)


async def generate_risk_register(state: dict) -> str:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Risk Register"

    headers = [
        "Risk ID",
        "Description",
        "Probability",
        "Impact",
        "Risk Score",
        "Owner",
        "Controls",
        "Control Test Result",
        "Consolidated Ref",
    ]
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = Font(bold=True)

    interim_map, fieldwork_map = _build_consolidated_maps(state)
    risk_matrix = state.get("risk_control_matrix", []) or []
    for idx, risk in enumerate(state.get("risk_entries", []), 2):
        finding_index = risk.get("finding_stage_index") or risk.get("finding_index")
        if not finding_index:
            finding_index = _find_finding_index(state, risk.get("description", ""))
        finding_stage = risk.get("finding_stage") or state.get("stage") or "BOTH"
        consolidated = (
            interim_map.get(finding_index) if finding_stage == "INTERIM" else fieldwork_map.get(finding_index)
        ) or {}
        consolidated_ref = f"CF-{finding_index:03d}" if consolidated and finding_index else None

        control_result = ""
        related_controls = risk.get("related_controls", []) or []
        for control in related_controls:
            match = next(
                (
                    entry
                    for entry in risk_matrix
                    if control.lower() in str(entry.get("control_description", "")).lower()
                ),
                None,
            )
            if match:
                control_result = match.get("test_result") or ""
                break
        ws.cell(row=idx, column=1, value=risk.get("id", f"RISK-{idx}"))
        ws.cell(row=idx, column=2, value=risk.get("description", ""))
        ws.cell(row=idx, column=3, value=risk.get("probability", 0))
        ws.cell(row=idx, column=4, value=risk.get("impact", 0))
        ws.cell(row=idx, column=5, value=risk.get("risk_score", 0))
        ws.cell(row=idx, column=6, value=risk.get("owner", ""))
        ws.cell(row=idx, column=7, value=str(risk.get("related_controls", [])))
        ws.cell(row=idx, column=8, value=control_result)
        ws.cell(row=idx, column=9, value=consolidated_ref)

    session_id = state["session_id"]
    path = OUTPUT_DIR / f"risk_register_{session_id}.xlsx"
    wb.save(path)
    return str(path)


async def generate_memo(state: dict) -> str:
    llm = get_llm_client("generation")
    anomaly_rules = {flag.get("rule") for flag in state["anomaly_flags"] if flag.get("rule")}
    anomaly_summary = f"{len(state['anomaly_flags'])} anomalies flagged: {', '.join(sorted(anomaly_rules))}"

    findings = _annotate_findings(state)
    file_names = _file_name_map(state)
    memo_findings = []
    for finding in findings:
        evidence_links = finding.get("evidence_links") or []
        evidence_reference = _render_evidence_links(evidence_links, file_names)
        if not evidence_reference:
            evidence_reference = finding.get("source_reference") or ""
        memo_findings.append(
            {
                "finding_id": finding.get("finding_id"),
                "stage": finding.get("stage"),
                "severity": finding.get("severity", "MEDIUM"),
                "confidence_score": finding.get("confidence_score", 0.0),
                "description": finding.get("description", ""),
                "evidence_reference": evidence_reference,
            }
        )

    prompt = MEMO_GENERATION_PROMPT.format(
        session_id=state["session_id"],
        stage=state["stage"],
        findings_json=json.dumps(memo_findings[:10], ensure_ascii=False),
        anomaly_summary=anomaly_summary,
    )
    memo_text = await llm(prompt)

    doc = Document()
    doc.add_heading(f"AUDIT MEMO - {state['stage']}", 0)
    doc.add_paragraph(f"Session ID: {state['session_id']}")
    doc.add_paragraph("")

    for line in memo_text.split("\n"):
        line = line.strip()
        if not line:
            continue
        if line.startswith(("1.", "2.", "3.", "4.", "5.")):
            doc.add_heading(line, level=1)
        else:
            doc.add_paragraph(line)

    memo_text_lower = memo_text.lower()
    missing_ids = [
        f["finding_id"] for f in memo_findings
        if f.get("finding_id") and f["finding_id"].lower() not in memo_text_lower
    ]
    if missing_ids:
        doc.add_heading("DANH SACH PHAT HIEN (ID CHINH XAC)", level=1)
        for finding in memo_findings:
            doc.add_paragraph(
                f"{finding['finding_id']}: {finding['description']}"
            )

    interim_map, fieldwork_map = _build_consolidated_maps(state)
    review_findings = []
    for idx, finding in enumerate(findings, 1):
        stage = finding.get("stage") or state.get("stage") or "BOTH"
        stage_index = int(finding.get("_stage_index") or idx)
        consolidated = (
            interim_map.get(stage_index) if stage == "INTERIM" else fieldwork_map.get(stage_index)
        ) or {}
        if consolidated.get("review_flag") or finding.get("manual_review_required"):
            review_findings.append(
                {
                    "id": finding.get("finding_id") or f"FND-{state['session_id'][:4].upper()}-{idx:03d}",
                    "description": finding.get("description", ""),
                    "confidence": finding.get("confidence_score", 0.0),
                }
            )

    if review_findings:
        doc.add_heading("MANUAL REVIEW REQUIRED", level=1)
        for finding in review_findings:
            doc.add_paragraph(
                f"{finding['id']}: {finding['description']} (confidence {finding['confidence']:.0%})"
            )

    path = OUTPUT_DIR / f"audit_memo_{state['session_id']}.docx"
    doc.save(path)
    return str(path)


async def generate_evidence_pdf(state: dict) -> str:
    session_id = state["session_id"]
    path = OUTPUT_DIR / f"evidence_{session_id}.pdf"

    if canvas is None or letter is None:
        lines = [f"Evidence Bundle - Session {session_id}", ""]
        for idx, finding in enumerate(_annotate_findings(state), 1):
            finding_id = finding.get("finding_id") or f"FND-{session_id[:4].upper()}-{idx:03d}"
            lines.append(f"{finding_id}: {finding.get('description', '')}")
        path.write_text("\n".join(lines), encoding="utf-8")
        return str(path)

    findings = _annotate_findings(state)
    file_names = _file_name_map(state)

    pdf = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    y = height - 50

    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, f"Evidence Bundle - Session {session_id}")
    y -= 30
    pdf.setFont("Helvetica", 10)

    for idx, finding in enumerate(findings, 1):
        if y < 100:
            pdf.showPage()
            y = height - 50
            pdf.setFont("Helvetica", 10)

        finding_id = finding.get("finding_id") or f"FND-{session_id[:4].upper()}-{idx:03d}"
        pdf.setFont("Helvetica-Bold", 11)
        pdf.drawString(50, y, f"{finding_id}: {finding.get('description', '')}")
        y -= 16

        pdf.setFont("Helvetica", 9)
        evidence_links = finding.get("evidence_links") or []
        if evidence_links:
            for link in evidence_links:
                reference = link.get("reference") or "Unknown"
                filename = file_names.get(link.get("source_file_id"), "")
                prefix = f"{filename}: " if filename else ""
                pdf.drawString(60, y, f"- {prefix}{reference}")
                y -= 12
        else:
            pdf.drawString(60, y, "- No evidence links available")
            y -= 12

        y -= 6

    pdf.save()
    return str(path)


async def generate_versioned_notes(state: dict) -> str:
    session_id = state["session_id"]
    path = OUTPUT_DIR / f"versioned_notes_{session_id}.md"
    timestamp = datetime.now(timezone.utc).isoformat()

    lines = [f"# Versioned Notes - {session_id}", "", f"Generated at: {timestamp}", ""]
    for entry in state.get("versioned_notes", []) or []:
        lines.append(f"- {entry}")

    path.write_text("\n".join(lines), encoding="utf-8")
    return str(path)
