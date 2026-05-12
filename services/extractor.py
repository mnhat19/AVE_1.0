from pathlib import Path

import pandas as pd
import pdfplumber
from docx import Document
from services.email_parser import parse_eml, parse_pst
from services.ocr_service import extract_text


def extract_pdf(path: str) -> dict:
    text_parts: list[str] = []
    tables: list[dict] = []
    try:
        with pdfplumber.open(path) as pdf:
            for page_index, page in enumerate(pdf.pages, start=1):
                text_parts.append(page.extract_text() or "")
                page_tables = page.extract_tables() or []
                for table_index, table in enumerate(page_tables, start=1):
                    if not table:
                        continue
                    tables.append({"page": page_index, "index": table_index, "rows": table})
            page_count = len(pdf.pages)
    except Exception as exc:
        return {"type": "text", "content": "", "error": f"pdf_parse_error:{exc}"}
    return {
        "type": "text",
        "content": "\n".join(text_parts),
        "pages": page_count,
        "tables": tables,
    }


def extract_docx(path: str) -> dict:
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        tables.append(rows)
    return {"type": "text", "content": "\n".join(paragraphs), "tables": tables}


def extract_excel(path: str) -> dict:
    xl = pd.ExcelFile(path)
    sheets = {}
    for sheet in xl.sheet_names:
        df = xl.parse(sheet)
        sheets[sheet] = df.to_dict(orient="records")
    return {"type": "table", "content": sheets, "sheet_names": xl.sheet_names}


def extract_csv(path: str) -> dict:
    df = pd.read_csv(path, encoding="utf-8-sig")
    return {"type": "table", "content": df.to_dict(orient="records"), "columns": list(df.columns)}


def extract_text(path: str) -> dict:
    content = Path(path).read_text(encoding="utf-8", errors="ignore")
    return {"type": "text", "content": content}


def extract_image(path: str) -> dict:
    result = extract_text(path)
    payload = {"type": "text", "content": result.get("text", "")}
    if result.get("error"):
        payload["error"] = result["error"]
    return payload


def extract_eml(path: str) -> dict:
    parsed = parse_eml(path)
    return {
        "type": "text",
        "content": parsed.get("content", ""),
        "metadata": parsed.get("metadata", {}),
    }


def extract_pst(path: str) -> dict:
    parsed = parse_pst(path)
    payload = {
        "type": "text",
        "content": parsed.get("content", ""),
        "metadata": parsed.get("metadata", {}),
    }
    if parsed.get("error"):
        payload["error"] = parsed["error"]
    return payload


EXTRACTORS = {
    "PDF": extract_pdf,
    "DOCX": extract_docx,
    "XLSX": extract_excel,
    "CSV": extract_csv,
    "TXT": extract_text,
    "IMAGE": extract_image,
    "EML": extract_eml,
    "PST": extract_pst,
}


def extract(file_record) -> dict:
    fn = EXTRACTORS.get(file_record.format)
    if not fn:
        return {"type": "text", "content": "", "error": "unsupported format"}
    return fn(file_record.file_path)
