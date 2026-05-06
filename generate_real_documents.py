from pathlib import Path
import csv
import random
from datetime import datetime, timedelta

import openpyxl
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

BASE_DIR = Path("d:/Document/AiB/AVE_1.2/sample_audit_documents")

def create_journal_csv(path: Path) -> None:
    rows = [["Date", "Reference", "AccountCode", "AccountName", "Debit", "Credit", "Description", "Preparer", "Approver"]]
    
    start_date = datetime(2026, 1, 1)
    
    for i in range(1, 501):
        d = start_date + timedelta(days=random.randint(0, 360))
        date_str = d.strftime("%Y-%m-%d")
        ref = f"JV-{i:05d}"
        
        is_weekend = d.weekday() >= 5
        is_rounded = random.random() < 0.05
        is_manual = random.random() < 0.05
        
        if is_rounded:
            amount = float(random.randint(1, 100) * 1000000)
        else:
            amount = round(random.uniform(1000, 500000), 2)
            
        desc = "Routine expense accrual"
        preparer = "Staff_A"
        approver = "Mgr_B"
        
        if is_manual:
            desc = "Manual override adjustment"
            approver = "" 
            
        if is_weekend:
            desc = "Urgent weekend adjustment"
            
        rows.append([date_str, ref, "642", "Admin Expense", amount, 0, desc, preparer, approver])
        rows.append([date_str, ref, "331", "Accounts Payable", 0, amount, desc, preparer, approver])
        
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)

def create_trial_balance(path: Path) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Trial Balance"
    headers = ["Account Code", "Account Name", "Opening Balance", "Debit", "Credit", "Closing Balance"]
    ws.append(headers)
    
    accounts = [
        ("111", "Cash", 50000, 150000, 120000, 80000),
        ("131", "Accounts Receivable", 100000, 50000, 40000, 110000),
        ("211", "Fixed Assets", 500000, 0, 0, 500000),
        ("331", "Accounts Payable", -80000, -20000, -60000, -160000),
        ("411", "Equity", -570000, 0, 0, -570000),
        ("511", "Revenue", 0, 0, -300000, -300000),
        ("642", "Admin Expense", 0, 340000, 0, 340000),
    ]
    
    for acc in accounts:
        ws.append(acc)
        
    wb.save(path)

def create_procurement_doc(path: Path) -> None:
    doc = Document()
    doc.add_heading("Standard Operating Procedure: Procurement", 0)
    
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph("This document outlines the standard operating procedure for procurement at the company, ensuring all purchases are authorized, valid, and properly recorded.")
    
    doc.add_heading("2. Vendor Selection", level=1)
    doc.add_paragraph("All new vendors must be approved by the Procurement Manager. A background check and conflict of interest declaration are required.")
    
    doc.add_heading("3. Purchase Requisitions & Orders", level=1)
    doc.add_paragraph("All purchases require a Purchase Requisition (PR).")
    doc.add_paragraph("Purchases below 10,000,000 VND can be approved by the Department Head.")
    doc.add_paragraph("Purchases above 10,000,000 VND require dual approval from the Department Head and the CFO.")
    doc.add_paragraph("However, urgent purchases (classified by Department Head) bypass CFO approval to avoid delays.") 
    
    doc.add_heading("4. Invoice Processing", level=1)
    doc.add_paragraph("The finance team performs a 3-way match (PO, Goods Receipt, Invoice) before payment.")
    
    doc.save(path)

def create_walkthrough_pdf(path: Path) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter)
    styles = getSampleStyleSheet()
    Story = []
    
    Story.append(Paragraph("Walkthrough Notes: Procurement Cycle", styles['Title']))
    Story.append(Spacer(1, 12))
    Story.append(Paragraph("Date of Walkthrough: Oct 15, 2026", styles['Normal']))
    Story.append(Paragraph("Auditor: John Doe", styles['Normal']))
    Story.append(Paragraph("Interviewee: Jane Smith (Procurement Manager)", styles['Normal']))
    Story.append(Spacer(1, 12))
    
    Story.append(Paragraph("Observation 1: Vendor Onboarding", styles['Heading2']))
    Story.append(Paragraph("Walked through the onboarding of 'Tech Supplies Inc'. The background check was present. However, the conflict of interest declaration was signed digitally without identity verification.", styles['Normal']))
    
    Story.append(Spacer(1, 12))
    Story.append(Paragraph("Observation 2: Purchase Order Approval", styles['Heading2']))
    Story.append(Paragraph("Selected a sample of 5 POs. 4 POs followed the standard workflow. 1 PO (PO-9921) for 50,000,000 VND was marked as 'Urgent' and bypassed CFO approval. Jane stated this is common practice at month-end.", styles['Normal'])) 
    
    Story.append(Spacer(1, 12))
    Story.append(Paragraph("Observation 3: 3-Way Match", styles['Heading2']))
    Story.append(Paragraph("Observed the AP clerk matching an invoice to a PO and GRN. The system allowed payment even though the GRN quantity was slightly less than the invoice quantity (within a 5% tolerance set in the system).", styles['Normal']))
    
    doc.build(Story)

def main() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    create_journal_csv(BASE_DIR / "journal_entries.csv")
    create_trial_balance(BASE_DIR / "trial_balance.xlsx")
    create_procurement_doc(BASE_DIR / "sop_procurement.docx")
    create_walkthrough_pdf(BASE_DIR / "walkthrough_notes.pdf")
    print("Real complete documents generated in sample_audit_documents/")

if __name__ == "__main__":
    main()
