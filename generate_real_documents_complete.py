from __future__ import annotations

from datetime import datetime, timedelta
from email.message import EmailMessage
from pathlib import Path
import csv
import random

import openpyxl
from openpyxl.chart import BarChart, LineChart, Reference
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from docx import Document
from PIL import Image, ImageDraw
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.shapes import Drawing, String
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

BASE_DIR = Path(__file__).resolve().parent / "sample_audit_documents"


def ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def write_csv(path: Path, rows: list[list[object]]) -> None:
    ensure_parent(path)
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def write_text(path: Path, content: str) -> None:
    ensure_parent(path)
    path.write_text(content, encoding="utf-8")


def styled_header(ws, header_row: int = 1) -> None:
    fill = PatternFill("solid", fgColor="D9E1F2")
    font = Font(bold=True)
    align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    for cell in ws[header_row]:
        cell.fill = fill
        cell.font = font
        cell.alignment = align


def autosize_columns(ws) -> None:
    for col_cells in ws.columns:
        max_len = 0
        for cell in col_cells:
            if cell.value is None:
                continue
            max_len = max(max_len, len(str(cell.value)))
        ws.column_dimensions[col_cells[0].column_letter].width = max(10, min(45, max_len + 2))


def build_pdf_report(
    path: Path,
    title: str,
    meta_lines: list[str],
    paragraphs: list[str],
    table_rows: list[list[str]] | None = None,
    footer: str | None = None,
) -> None:
    ensure_parent(path)
    doc = SimpleDocTemplate(str(path), pagesize=letter, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph(title, styles["Title"]))
    story.append(Spacer(1, 10))

    for line in meta_lines:
        story.append(Paragraph(line, styles["Normal"]))
    story.append(Spacer(1, 12))

    for para in paragraphs:
        story.append(Paragraph(para, styles["Normal"]))
        story.append(Spacer(1, 8))

    if table_rows:
        table = Table(table_rows, hAlign="LEFT")
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                    ("FONT", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ]
            )
        )
        story.append(Spacer(1, 10))
        story.append(table)

    if footer:
        story.append(Spacer(1, 14))
        story.append(Paragraph(footer, styles["Normal"]))

    doc.build(story)


def create_cover_pdf(
    path: Path,
    title: str,
    subtitle: str,
    client: str,
    period: str,
    prepared_by: str,
    prepared_date: str,
) -> None:
    ensure_parent(path)
    pdf = canvas.Canvas(str(path), pagesize=letter)
    width, height = letter
    pdf.setFillColor(colors.HexColor("#1F4E79"))
    pdf.rect(0, height - 140, width, 140, stroke=0, fill=1)
    pdf.setFillColor(colors.white)
    pdf.setFont("Helvetica-Bold", 22)
    pdf.drawString(60, height - 80, title)
    pdf.setFont("Helvetica", 12)
    pdf.drawString(60, height - 105, subtitle)

    pdf.setFillColor(colors.black)
    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(60, height - 190, "Client")
    pdf.setFont("Helvetica", 12)
    pdf.drawString(60, height - 210, client)

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(60, height - 250, "Period")
    pdf.setFont("Helvetica", 12)
    pdf.drawString(60, height - 270, period)

    pdf.setFont("Helvetica-Bold", 14)
    pdf.drawString(60, height - 310, "Prepared By")
    pdf.setFont("Helvetica", 12)
    pdf.drawString(60, height - 330, prepared_by)
    pdf.drawString(60, height - 350, prepared_date)

    pdf.setFont("Helvetica", 10)
    pdf.setFillColor(colors.HexColor("#555555"))
    pdf.drawString(60, 60, "Confidential - For audit engagement use only")
    pdf.save()


def create_management_dashboard_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Dashboard"
    sheet.append(["Metric", "Value"])
    rows = [
        ("Controls Tested", 24),
        ("Exceptions", 3),
        ("High Risk Findings", 2),
        ("Open Items", 5),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)

    chart = BarChart()
    chart.title = "Audit KPI Summary"
    chart.y_axis.title = "Count"
    chart.x_axis.title = "Metric"
    data = Reference(sheet, min_col=2, min_row=1, max_row=1 + len(rows))
    cats = Reference(sheet, min_col=1, min_row=2, max_row=1 + len(rows))
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 8
    chart.width = 14
    sheet.add_chart(chart, "D2")

    autosize_columns(sheet)
    workbook.save(path)


def create_cash_flow_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Cash Flow"
    sheet.append(["Month", "Operating", "Investing", "Financing"])
    rows = [
        ("Jan", 320000, -45000, 80000),
        ("Feb", 280000, -60000, 50000),
        ("Mar", 350000, -75000, 60000),
        ("Apr", 290000, -50000, 40000),
        ("May", 410000, -82000, 70000),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)

    chart = LineChart()
    chart.title = "Cash Flow Trend"
    chart.y_axis.title = "VND"
    data = Reference(sheet, min_col=2, min_row=1, max_col=4, max_row=1 + len(rows))
    cats = Reference(sheet, min_col=1, min_row=2, max_row=1 + len(rows))
    chart.add_data(data, titles_from_data=True)
    chart.set_categories(cats)
    chart.height = 9
    chart.width = 15
    sheet.add_chart(chart, "E2")

    autosize_columns(sheet)
    workbook.save(path)


def create_kpi_summary_pdf(path: Path) -> None:
    ensure_parent(path)
    doc = SimpleDocTemplate(str(path), pagesize=letter, topMargin=0.7 * inch, bottomMargin=0.7 * inch)
    styles = getSampleStyleSheet()
    story = []
    story.append(Paragraph("Audit KPI Summary", styles["Title"]))
    story.append(Spacer(1, 10))
    story.append(Paragraph("Engagement: FY2026 Interim + Fieldwork", styles["Normal"]))
    story.append(Paragraph("Prepared by: Audit Analytics Team", styles["Normal"]))
    story.append(Spacer(1, 14))

    drawing = Drawing(420, 220)
    chart = VerticalBarChart()
    chart.x = 40
    chart.y = 30
    chart.height = 150
    chart.width = 340
    chart.data = [[85, 72, 64, 90]]
    chart.categoryAxis.categoryNames = ["Controls", "Evidence", "Reconciliations", "Exceptions"]
    chart.valueAxis.valueMin = 0
    chart.valueAxis.valueMax = 100
    chart.valueAxis.valueStep = 20
    chart.bars[0].fillColor = colors.HexColor("#4F81BD")
    drawing.add(chart)
    drawing.add(String(40, 190, "Completion (%)", fontSize=10))
    story.append(drawing)

    story.append(Spacer(1, 12))
    story.append(Paragraph("Summary: 2 high-risk findings remain open pending client response.", styles["Normal"]))
    doc.build(story)


def create_reconciliation_summary_pdf(path: Path) -> None:
    build_pdf_report(
        path,
        "Bank Reconciliation Summary",
        [
            "Period: February 2026",
            "Prepared by: Treasury Ops",
            "Reviewed by: Finance Manager",
        ],
        [
            "The reconciliation identified one outstanding transfer requiring follow-up.",
            "All other reconciling items cleared within the expected window.",
        ],
        table_rows=[
            ["Item", "Ledger", "Bank", "Difference", "Status"],
            ["Cash balance", "11,850,000", "11,600,000", "250,000", "Open"],
            ["Outstanding deposit", "0", "0", "0", "Matched"],
            ["Unpresented cheques", "0", "0", "0", "Matched"],
        ],
        footer="Reconciliation prepared using bank statement BA-2026-02.",
    )


def create_interim_procurement_docx(path: Path) -> None:
    ensure_parent(path)
    doc = Document()
    doc.add_heading("Standard Operating Procedure: Procurement", 0)
    doc.add_paragraph(
        "This procedure governs purchase requisitions, vendor onboarding, invoice approval, and payment release."
    )

    doc.add_heading("Approval Thresholds", level=1)
    table = doc.add_table(rows=1, cols=4)
    headers = ["Threshold", "Approver", "Evidence", "Notes"]
    for idx, header in enumerate(headers):
        table.cell(0, idx).text = header
    rows = [
        ["<= VND 10,000,000", "Department Head", "Approved PR", "Single approval"],
        ["> VND 10,000,000", "Dept Head + CFO", "Dual approval", "Two-level sign-off"],
        ["Urgent purchases", "Dept Head", "Urgency memo", "CFO sign-off within 2 business days"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value

    doc.add_heading("Control Activities", level=1)
    doc.add_paragraph("Three-way match between PO, goods receipt, and supplier invoice is mandatory before payment.")
    doc.add_paragraph("Vendor master changes must be reviewed monthly with documented approval.")
    doc.add_paragraph("Manual journal entries above VND 1,000,000 require reviewer sign-off and rationale.")

    doc.add_heading("Sign-Off", level=1)
    doc.add_paragraph("Prepared by: Procurement Manager | Date: 2026-10-10")
    doc.add_paragraph("Reviewed by: Finance Director | Date: 2026-10-12")
    doc.save(path)


def create_internal_policy_docx(path: Path) -> None:
    ensure_parent(path)
    doc = Document()
    doc.add_heading("Internal Control Policy", 0)
    doc.add_paragraph("All procurement and finance transactions must retain supporting evidence for audit traceability.")
    doc.add_paragraph("System overrides must be logged and reviewed weekly by the control owner.")
    doc.add_paragraph("Exceptions exceeding tolerance thresholds require escalation to the risk committee.")
    doc.add_paragraph("Document retention period: 7 years for finance records, 10 years for contracts.")
    doc.save(path)


def create_walkthrough_pdf(path: Path) -> None:
    build_pdf_report(
        path,
        "Walkthrough Notes: Procurement Cycle",
        [
            "Date: 2026-10-15",
            "Auditor: John Doe",
            "Interviewee: Jane Smith, Procurement Manager",
        ],
        [
            "Observation 1: Conflict of interest declarations were signed digitally without identity verification.",
            "Observation 2: One high-value purchase order was marked urgent and bypassed CFO approval at month-end.",
            "Observation 3: Payment was processed when GRN quantity differed slightly from invoice quantity within tolerance.",
        ],
        table_rows=[
            ["Sample", "PO", "Amount (VND)", "Result"],
            ["1", "PO-9921", "50,000,000", "Urgent override"],
            ["2", "PO-9933", "18,500,000", "Compliant"],
            ["3", "PO-9941", "9,800,000", "Compliant"],
        ],
        footer="Prepared by audit team. Evidence retained in working paper WP-INT-03.",
    )


def create_risk_matrix_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Risk Matrix"
    sheet.append(["risk_id", "risk", "control", "test_procedure", "owner", "severity"])
    rows = [
        ("IM-001", "Unauthorized procurement approval", "Dual approval required above threshold", "Inspect approval evidence for selected samples", "Procurement Manager", "High"),
        ("IM-002", "Incomplete walkthrough evidence", "Walkthroughs documented and signed", "Trace walkthrough notes to signed minutes", "Finance Manager", "Medium"),
        ("IM-003", "Supplier master changes without review", "Monthly vendor master review", "Reperform review of vendor change log", "AP Manager", "High"),
        ("IM-004", "Manual journal overrides", "Reviewer sign-off required", "Inspect approval for manual journals", "Controller", "Medium"),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)
    autosize_columns(sheet)
    workbook.save(path)


def create_test_of_controls_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "ToC"
    sheet.append(["test_id", "control", "sample", "result", "remark", "evidence_ref"])
    rows = [
        ("TOC-01", "PO approval threshold", "5 purchases", "PASS", "One urgent override noted", "WP-INT-01"),
        ("TOC-02", "3-way match", "4 invoices", "PASS", "Minor quantity variance within tolerance", "WP-INT-02"),
        ("TOC-03", "Vendor master review", "Monthly review", "FAIL", "Evidence of review was not dated", "WP-INT-04"),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)
    autosize_columns(sheet)
    workbook.save(path)


def create_control_questionnaire_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Control Questionnaire"
    sheet.append(["question_id", "question", "response", "owner", "evidence"])
    rows = [
        ("CQ-01", "Are vendor approvals documented?", "Yes", "Procurement", "Vendor_Approval_Log.pdf"),
        ("CQ-02", "Is a 3-way match performed before payment?", "Yes", "AP", "3-Way_Match_Report.xlsx"),
        ("CQ-03", "Are overrides monitored weekly?", "Partially", "Finance", "Override_Log.csv"),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)
    autosize_columns(sheet)
    workbook.save(path)


def create_eml(path: Path, subject: str, sender: str, recipient: str, body: str) -> None:
    ensure_parent(path)
    message = EmailMessage()
    message["Subject"] = subject
    message["From"] = sender
    message["To"] = recipient
    message.set_content(body)
    path.write_bytes(message.as_bytes())


def create_journal_csv(path: Path) -> None:
    random.seed(42)
    rows = [["date", "reference", "account_code", "account_name", "debit", "credit", "description", "prepared_by", "approved_by", "cost_center"]]
    start_date = datetime(2026, 1, 1)

    for i in range(1, 151):
        current_date = start_date + timedelta(days=random.randint(0, 360))
        reference = f"JV-{i:05d}"
        weekend = current_date.weekday() >= 5
        rounded = random.random() < 0.08
        manual = random.random() < 0.06
        amount = float(random.randint(1, 80) * 100000) if rounded else round(random.uniform(1250, 685000), 2)
        description = "Routine expense accrual"
        prepared_by = "AP_Clerk"
        approved_by = "AP_Manager"
        cost_center = random.choice(["HQ", "Sales", "Ops", "IT"])

        if weekend:
            description = "Weekend adjustment posted after close"
        if manual:
            description = "Manual override adjustment"
            approved_by = ""

        rows.append([current_date.strftime("%Y-%m-%d"), reference, "642", "Administrative Expense", amount, 0, description, prepared_by, approved_by, cost_center])
        rows.append([current_date.strftime("%Y-%m-%d"), reference, "331", "Accounts Payable", 0, amount, description, prepared_by, approved_by, cost_center])

    write_csv(path, rows)


def create_gl_export_csv(path: Path) -> None:
    rows = [["date", "doc_no", "account", "description", "debit", "credit", "currency", "counterparty"]]
    start_date = datetime(2026, 1, 1)
    for i in range(1, 51):
        date_str = (start_date + timedelta(days=i * 3)).strftime("%Y-%m-%d")
        rows.append([date_str, f"GL-{i:04d}", "111", "Cash receipt", 250000 + i * 2500, 0, "VND", "Customer A"])
        rows.append([date_str, f"GL-{i:04d}", "511", "Revenue", 0, 250000 + i * 2500, "VND", "Customer A"])
    write_csv(path, rows)


def create_trial_balance_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "TB"
    sheet.append(["account_code", "account_name", "debit", "credit", "balance"])
    rows = [
        ("111", "Cash", 11850000, 0, 11850000),
        ("131", "Accounts Receivable", 6120000, 0, 6120000),
        ("141", "Prepayments", 980000, 0, 980000),
        ("211", "Fixed Assets", 18200000, 0, 18200000),
        ("331", "Accounts Payable", 0, 7425000, -7425000),
        ("341", "Payroll Liabilities", 0, 1885000, -1885000),
        ("411", "Equity", 0, 22000000, -22000000),
        ("511", "Revenue", 0, 15400000, -15400000),
        ("642", "Administrative Expense", 4250000, 0, 4250000),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)
    autosize_columns(sheet)
    workbook.save(path)


def create_lead_schedule_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Lead Schedule"
    sheet.append(["account_code", "account_name", "tb_balance", "lead_adjustment", "final_balance"])
    rows = [
        ("111", "Cash", 11850000, 0, 11850000),
        ("131", "Accounts Receivable", 6120000, -120000, 6000000),
        ("211", "Fixed Assets", 18200000, 0, 18200000),
        ("331", "Accounts Payable", -7425000, 0, -7425000),
        ("511", "Revenue", -15400000, 0, -15400000),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)
    autosize_columns(sheet)
    workbook.save(path)


def create_ageing_csv(path: Path) -> None:
    rows = [
        ["customer", "current", "31_60", "61_90", "over_90", "total"],
        ["Alpha Trading", 420000, 180000, 0, 0, 600000],
        ["Beta Logistics", 180000, 240000, 120000, 0, 540000],
        ["Gamma Retail", 0, 160000, 210000, 90000, 460000],
    ]
    write_csv(path, rows)


def create_reconciliation_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "Reconciliation"
    sheet.append(["item", "ledger", "bank", "difference", "status", "note"])
    rows = [
        ("Cash balance", 11850000, 11600000, 250000, "Open", "Transfer pending investigation"),
        ("Outstanding deposit", 0, 0, 0, "Matched", "Cleared in March"),
        ("Unpresented cheques", 0, 0, 0, "Matched", "No exceptions"),
    ]
    for row in rows:
        sheet.append(row)
    styled_header(sheet)
    autosize_columns(sheet)
    workbook.save(path)


def create_bank_statement_pdf(path: Path) -> None:
    build_pdf_report(
        path,
        "Bank Statement - February 2026",
        [
            "Account: 123-456-789",
            "Period: 2026-02-01 to 2026-02-28",
            "Prepared by: Treasury Operations",
        ],
        [
            "Opening balance: VND 4,820,000,000",
            "Closing balance: VND 4,312,500,000",
            "One manual transfer reference lacks a full remittance narrative and requires follow-up.",
        ],
        table_rows=[
            ["Date", "Description", "Debit", "Credit", "Balance"],
            ["2026-02-05", "Supplier settlement", "85,000,000", "", "4,735,000,000"],
            ["2026-02-15", "Payroll", "120,000,000", "", "4,615,000,000"],
            ["2026-02-20", "Customer receipt", "", "250,000,000", "4,865,000,000"],
        ],
        footer="Statement generated by the bank core system. Confirm with bank if discrepancies are noted.",
    )


def create_inventory_image(path: Path) -> None:
    ensure_parent(path)
    image = Image.new("RGB", (1200, 800), "white")
    draw = ImageDraw.Draw(image)
    draw.rectangle((40, 40, 1160, 760), outline="black", width=4)
    draw.text((80, 70), "Physical Inventory Count", fill="black")
    draw.text((80, 120), "Location: Warehouse A", fill="black")
    draw.text((80, 160), "Count date: 2026-12-28", fill="black")
    draw.text((80, 210), "Item: Finished goods", fill="black")
    draw.text((80, 250), "Quantity: 4,250 cartons", fill="black")
    draw.text((80, 290), "Variance noted: 2 damaged cartons pending review", fill="black")
    draw.text((80, 360), "Signed by storekeeper and supervisor", fill="black")
    image.save(path)


def create_contract_docx(path: Path) -> None:
    ensure_parent(path)
    doc = Document()
    doc.add_heading("Vendor Supply Agreement", 0)
    doc.add_paragraph("This agreement sets out pricing, delivery, warranty, and payment terms for recurring supplies.")
    doc.add_paragraph("Delivery confirmations must be matched to the purchase order before invoice approval.")
    doc.add_paragraph("Any deviation in quantity greater than 5 percent must be escalated to procurement management.")
    doc.add_paragraph("Effective period: 2026-01-01 to 2027-12-31")
    doc.save(path)


def create_corrupted_pdf(path: Path) -> None:
    ensure_parent(path)
    path.write_bytes(b"%PDF-1.4\n% corrupted sample file\nthis is not a complete pdf body")


def create_trial_balance_missing_columns_xlsx(path: Path) -> None:
    workbook = openpyxl.Workbook()
    sheet = workbook.active
    sheet.title = "TB"
    sheet.append(["account_code", "account_name", "debit", "credit"])
    sheet.append(["111", "Cash", 10000000, 0])
    sheet.append(["331", "Accounts Payable", 0, 2500000])
    workbook.save(path)


def create_journal_with_errors_csv(path: Path) -> None:
    rows = [
        ["date", "reference", "debit", "credit", "description", "amount"],
        ["2026-02-01", "JV-9001", 1000000, 1000000, "Opening entry", 1000000],
        ["2026-02-02", "JV-9002", "not-a-number", 500000, "Malformed debit", 500000],
        ["", "JV-9003", 250000, 250000, "Missing date", 250000],
        ["2026-02-05", "JV-9004", 750000, 750000, "Valid control sample", 750000],
    ]
    write_csv(path, rows)


def create_pst_placeholder(path: Path) -> None:
    ensure_parent(path)
    path.write_bytes(
        b"PST placeholder for compatibility testing only.\n"
        b"A valid Outlook PST export is not generated in this repository."
    )


def create_large_text_file(path: Path, size_mb: int = 5) -> None:
    ensure_parent(path)
    chunk = ("Oversize sample attachment for size limit testing.\n" * 2000).encode("utf-8")
    target = size_mb * 1024 * 1024
    with path.open("wb") as handle:
        written = 0
        while written < target:
            handle.write(chunk)
            written += len(chunk)


def main() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)

    # Interim documents
    create_cover_pdf(
        BASE_DIR / "INTERIM" / "audit_planning_cover.pdf",
        "Audit Planning Package",
        "Interim Audit Documentation",
        "Blue Harbor Manufacturing Co.",
        "FY2026 (Interim)",
        "Audit Engagement Team",
        "Prepared: 2026-10-12",
    )
    create_interim_procurement_docx(BASE_DIR / "INTERIM" / "sop_procurement.docx")
    create_internal_policy_docx(BASE_DIR / "INTERIM" / "policy_internal.docx")
    create_walkthrough_pdf(BASE_DIR / "INTERIM" / "walkthrough_notes.pdf")
    create_risk_matrix_xlsx(BASE_DIR / "INTERIM" / "risk_matrix.xlsx")
    create_test_of_controls_xlsx(BASE_DIR / "INTERIM" / "test_of_controls.xlsx")
    create_control_questionnaire_xlsx(BASE_DIR / "INTERIM" / "control_questionnaire.xlsx")
    create_eml(
        BASE_DIR / "INTERIM" / "email_confirmation.eml",
        "Procurement walkthrough confirmation",
        "jane.smith@company.vn",
        "audit.team@company.vn",
        "Please confirm the three-way match samples and the urgent purchase approval exception discussed during the walkthrough.",
    )

    # Fieldwork documents
    create_cover_pdf(
        BASE_DIR / "FIELDWORK" / "fieldwork_cover.pdf",
        "Fieldwork Evidence Package",
        "Year-End Audit Deliverables",
        "Blue Harbor Manufacturing Co.",
        "FY2026 (Fieldwork)",
        "Audit Engagement Team",
        "Prepared: 2026-12-30",
    )
    create_journal_csv(BASE_DIR / "FIELDWORK" / "journal_entries.csv")
    create_gl_export_csv(BASE_DIR / "FIELDWORK" / "gl_export.csv")
    create_trial_balance_xlsx(BASE_DIR / "FIELDWORK" / "trial_balance.xlsx")
    create_lead_schedule_xlsx(BASE_DIR / "FIELDWORK" / "lead_schedule.xlsx")
    create_ageing_csv(BASE_DIR / "FIELDWORK" / "ageing.csv")
    create_reconciliation_xlsx(BASE_DIR / "FIELDWORK" / "reconciliation.xlsx")
    create_management_dashboard_xlsx(BASE_DIR / "FIELDWORK" / "management_dashboard.xlsx")
    create_cash_flow_xlsx(BASE_DIR / "FIELDWORK" / "cash_flow_analysis.xlsx")
    create_bank_statement_pdf(BASE_DIR / "FIELDWORK" / "bank_statement.pdf")
    create_kpi_summary_pdf(BASE_DIR / "FIELDWORK" / "kpi_summary.pdf")
    create_reconciliation_summary_pdf(BASE_DIR / "FIELDWORK" / "reconciliation_summary.pdf")
    create_contract_docx(BASE_DIR / "FIELDWORK" / "contract.docx")
    create_inventory_image(BASE_DIR / "FIELDWORK" / "inventory_count.png")
    create_eml(
        BASE_DIR / "FIELDWORK" / "confirmation_letter.eml",
        "Bank confirmation follow-up",
        "treasury@bank.com",
        "audit.team@company.vn",
        "Attached is the balance confirmation for year-end audit procedures. Please review the outstanding items and respond by Friday.",
    )

    write_text(
        BASE_DIR / "FIELDWORK" / "Bank_Statement.txt",
        "Bank statement summary for February 2026 with one unmatched transfer reference.",
    )
    write_text(
        BASE_DIR / "FIELDWORK" / "Confirmation_Letter.txt",
        "Third-party confirmation request for year-end cash balance.",
    )
    write_text(
        BASE_DIR / "FIELDWORK" / "Contract.txt",
        "Supplier contract summary with quantity variance escalation clause.",
    )

    # Edge cases
    create_corrupted_pdf(BASE_DIR / "edge_cases" / "Corrupted_PDF.pdf")
    create_journal_with_errors_csv(BASE_DIR / "edge_cases" / "Journal_With_Errors.csv")
    create_trial_balance_missing_columns_xlsx(BASE_DIR / "edge_cases" / "Trial_Balance_Missing_Columns.xlsx")
    create_pst_placeholder(BASE_DIR / "edge_cases" / "Mailbox_Sample.pst")
    create_large_text_file(BASE_DIR / "edge_cases" / "Oversize_Attachment.txt", size_mb=5)

    print("Real complete documents generated in sample_audit_documents/")


if __name__ == "__main__":
    main()
