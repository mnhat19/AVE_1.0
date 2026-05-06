from pathlib import Path
import csv

import openpyxl
from docx import Document
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas

BASE_DIR = Path(__file__).resolve().parent


def create_journal_csv(path: Path) -> None:
    rows = [
        ["date", "reference", "debit", "credit", "description", "amount"],
        ["2026-01-02", "JV-0001", 1000000, 1000000, "Opening entry", 1000000],
        ["2026-01-05", "JV-0002", 500000, 500000, "Payroll accrual", 500000],
        ["2026-01-08", "JV-0003", 2500000, 2500000, "Inventory adjustment", 2500000],
        ["2026-01-11", "JV-0004", 750000, 750000, "Vendor invoice", 750000],
        ["2026-01-14", "JV-0005", 1500000, 1500000, "Revenue accrual", 1500000],
        ["2026-01-17", "JV-0006", 200000, 200000, "Small expense", 200000],
        ["2026-01-20", "JV-0007", 900000, 900000, "Utility accrual", 900000],
        ["2026-01-23", "JV-0008", 1200000, 1200000, "Manual adjustment", 1200000],
        ["2026-01-26", "JV-0009", 3000000, 3000000, "Year-end adjustment", 3000000],
        ["2026-01-29", "JV-0010", 650000, 650000, "Vendor accrual", 650000],
    ]
    with path.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerows(rows)


def create_trial_balance(path: Path) -> None:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TB"
    ws.append(["account_code", "account_name", "debit", "credit", "balance"])
    ws.append(["111", "Cash", 10000000, 0, 10000000])
    ws.append(["131", "Accounts Receivable", 5000000, 0, 5000000])
    ws.append(["211", "Fixed Assets", 15000000, 0, 15000000])
    ws.append(["331", "Accounts Payable", 0, 4000000, -4000000])
    ws.append(["411", "Equity", 0, 26000000, -26000000])
    wb.save(path)


def create_procurement_doc(path: Path) -> None:
    doc = Document()
    doc.add_heading("Procurement SOP", 0)
    doc.add_paragraph("Purchases require dual approval for invoices above 50,000,000 VND.")
    doc.add_paragraph("Monthly review of vendor master data is required.")
    doc.save(path)


def create_walkthrough_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=letter)
    c.setFont("Helvetica", 11)
    c.drawString(72, 720, "Walkthrough Notes")
    c.drawString(72, 700, "Control testing performed on sample purchase invoices.")
    c.drawString(72, 680, "No deviations noted in the sample reviewed.")
    c.showPage()
    c.save()


def main() -> None:
    BASE_DIR.mkdir(parents=True, exist_ok=True)
    create_journal_csv(BASE_DIR / "journal_entries.csv")
    create_trial_balance(BASE_DIR / "trial_balance.xlsx")
    create_procurement_doc(BASE_DIR / "sop_procurement.docx")
    create_walkthrough_pdf(BASE_DIR / "walkthrough_notes.pdf")
    print("Test data generated in test_data/")


if __name__ == "__main__":
    main()
